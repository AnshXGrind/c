"""
Resume Parser Service
Handles extraction of text from PDF and DOCX files.
"""

import os
import re
from typing import Optional
from io import BytesIO

import PyPDF2
from docx import Document


class ResumeParser:
    """
    Multi-format resume parser supporting PDF and DOCX files.
    Handles multi-page documents with robust text extraction.
    """
    
    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".doc"]
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a resume file.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.supported_formats}")
        
        if ext == ".pdf":
            return self._extract_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._extract_from_docx(file_path)
        
        return ""
    
    def extract_from_bytes(self, content: bytes, filename: str) -> str:
        """
        Extract text from file content bytes.
        
        Args:
            content: File content as bytes
            filename: Original filename (for extension detection)
            
        Returns:
            Extracted text content
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".pdf":
            return self._extract_pdf_from_bytes(content)
        elif ext in [".docx", ".doc"]:
            return self._extract_docx_from_bytes(content)
        
        raise ValueError(f"Unsupported format: {ext}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file using PyPDF2."""
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {str(e)}")
        
        return "\n\n".join(text_parts)
    
    def _extract_pdf_from_bytes(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        text_parts = []
        
        try:
            reader = PyPDF2.PdfReader(BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {str(e)}")
        
        return "\n\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        text_parts = []
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {str(e)}")
        
        return "\n".join(text_parts)
    
    def _extract_docx_from_bytes(self, content: bytes) -> str:
        """Extract text from DOCX bytes."""
        text_parts = []
        
        try:
            doc = Document(BytesIO(content))
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {str(e)}")
        
        return "\n".join(text_parts)
    
    def extract_sections(self, text: str) -> dict:
        """
        Attempt to extract common resume sections.
        
        Args:
            text: Resume text content
            
        Returns:
            Dictionary with section names as keys and content as values
        """
        sections = {}
        
        # Common section headers
        section_patterns = [
            (r"(?i)(education|academic)", "education"),
            (r"(?i)(experience|work\s*history|employment)", "experience"),
            (r"(?i)(skills|technical\s*skills|core\s*competencies)", "skills"),
            (r"(?i)(projects|portfolio)", "projects"),
            (r"(?i)(certifications?|certificates?)", "certifications"),
            (r"(?i)(summary|objective|profile)", "summary"),
            (r"(?i)(awards?|achievements?|honors?)", "awards"),
        ]
        
        lines = text.split("\n")
        current_section = "header"
        sections[current_section] = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if this line is a section header
            found_section = False
            for pattern, section_name in section_patterns:
                if re.match(pattern, line_stripped) and len(line_stripped) < 50:
                    current_section = section_name
                    if current_section not in sections:
                        sections[current_section] = []
                    found_section = True
                    break
            
            if not found_section and line_stripped:
                sections[current_section].append(line_stripped)
        
        # Join lines in each section
        return {k: "\n".join(v) for k, v in sections.items() if v}
